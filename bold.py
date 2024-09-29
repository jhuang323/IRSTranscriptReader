from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBoxHorizontal, LTChar
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from collections import Counter
import docx

#User Settings: Name of the script
TargetPDFFile = "Wage and Income Transcript.pdf"
OutPutDocFileName = "BoldWageIncTranscriptHeader"

def extract_and_count_bold_texts(pdf_path, font_size_threshold=16):
    texts = []

    # Create a resource manager
    rsrcmgr = PDFResourceManager()
    
    # Set up a device and interpreter
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    # Open the PDF file
    with open(pdf_path, "rb") as fp:
        for page in PDFPage.get_pages(fp):
            interpreter.process_page(page)
            layout = device.get_result()
            
            for element in layout:
                if isinstance(element, LTTextBoxHorizontal):
                    for text_line in element:
                        line_text = text_line.get_text().strip()
                        for char in text_line:
                            if isinstance(char, LTChar) and 'Bold' in char.fontname and char.size >= font_size_threshold:
                                texts.append(line_text)
                                break  # Break after adding the line text once

    # Count occurrences of each text segment
    counts = Counter(texts)

    return counts

def WordDocExport(Adict):
    TheWdDoc = docx.Document()

    TheWdDoc.add_heading("Summary of Documents in Transcript")

    for text, count in bold_text_counts.items():
        TheWdDoc.add_paragraph(f"{text} (Count: {count})")

    TheWdDoc.save(OutPutDocFileName + ".docx")
    


if __name__ == "__main__":
    pdf_path = r"Wage and Income Transcript.pdf"
    bold_text_counts = extract_and_count_bold_texts(pdf_path)

    WordDocExport(bold_text_counts.items())

    

