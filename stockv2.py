#Given an IRS wage and Income transcript this program will extract the stock amounts and categorize them as Short or Long term
import pdfplumber
import docx

#User Settings: Name of the script
TargetPDFFile = "Wage and Income Transcript.pdf"
OutPutDocFileName = "StockWageIncTranscriptSummary"

def extract_and_group_by_fin_and_gain_loss_code(pdf_path):
    data = {}
    FINOccuranceDict = dict()
    current_fin = None
    current_code = None

    with pdfplumber.open(pdf_path) as pdf:
        total_text = ""
        for page in pdf.pages:
            total_text += page.extract_text()

    lines = total_text.split("\n")

    for i, line in enumerate(lines):
        if "Payer's Federal Identification Number (FIN):" in line:
            current_code = None
            current_fin = line.split("Payer's Federal Identification Number (FIN):")[-1].strip()
            if current_fin not in data:
                data[current_fin] = {}
            if current_fin not in FINOccuranceDict:
                FINOccuranceDict[current_fin] = 1
            else:
                FINOccuranceDict[current_fin] += 1

        if "Type of Gain or Loss Code:" in line:
            current_code = line.split("Type of Gain or Loss Code:")[-1].strip()
            if current_fin and current_code not in data[current_fin]:
                data[current_fin][current_code] = {'Total Proceeds': 0, 'Total Cost or Basis': 0}

        if "Proceeds:" in line:
            value = line.split("Proceeds:")[-1].strip().replace(',', '').replace('$', '')
            try:
                proceeds_value = float(value)
                if current_fin and current_code:
                    if current_code not in data[current_fin]:
                        data[current_fin][current_code] = {'Total Proceeds': 0, 'Total Cost or Basis': 0}
                    data[current_fin][current_code]['Total Proceeds'] += proceeds_value
                else:
                    # Look for the next "Type of Gain or Loss Code" if not found before
                    for next_line in lines[i+1:]:
                        if "Type of Gain or Loss Code:" in next_line:
                            current_code = next_line.split("Type of Gain or Loss Code:")[-1].strip()
                            if current_fin and current_code not in data[current_fin]:
                                data[current_fin][current_code] = {'Total Proceeds': 0, 'Total Cost or Basis': 0}
                            data[current_fin][current_code]['Total Proceeds'] += proceeds_value
                            break
            except ValueError:
                pass

        if "Cost or Basis:" in line:
            value = line.split("Cost or Basis:")[-1].strip().replace(',', '').replace('$', '')
            try:
                cost_basis_value = float(value)
                if current_fin and current_code:
                    if current_code not in data[current_fin]:
                        data[current_fin][current_code] = {'Total Proceeds': 0, 'Total Cost or Basis': 0}
                    data[current_fin][current_code]['Total Cost or Basis'] += cost_basis_value
                else:
                    # Look for the next "Type of Gain or Loss Code" if not found before
                    for next_line in lines[i+1:]:
                        if "Type of Gain or Loss Code:" in next_line:
                            current_code = next_line.split("Type of Gain or Loss Code:")[-1].strip()
                            if current_fin and current_code not in data[current_fin]:
                                data[current_fin][current_code] = {'Total Proceeds': 0, 'Total Cost or Basis': 0}
                            data[current_fin][current_code]['Total Cost or Basis'] += cost_basis_value
                            break
            except ValueError:
                pass

    return (data,FINOccuranceDict)

    

def WordDocExport(Adict,FinalOccuranceDict):
    TheWdDoc = docx.Document()

    TheWdDoc.add_heading("Summary of Stock in Transcript")

    # Print the aggregated data
    for fin, fin_data in Adict.items():
        TheWdDoc.add_paragraph(f"Payer's Federal Identification Number (FIN): {fin}")
        for code, values in fin_data.items():
            TheWdDoc.add_paragraph(f"  Type of Gain or Loss Code: {code}")
            TheWdDoc.add_paragraph(f"    Total Proceeds: ${values['Total Proceeds']:.2f}")
            TheWdDoc.add_paragraph(f"    Total Cost or Basis: ${values['Total Cost or Basis']:.2f}")
        TheWdDoc.add_paragraph(f" Count: {FinalOccuranceDict[fin]}")
        TheWdDoc.add_paragraph("-" * 40)

    TheWdDoc.save(OutPutDocFileName + ".docx")

if __name__ == "__main__":
    pdf_path = r"Wage and Income Transcript.pdf"
    ThereturnDataTuple = extract_and_group_by_fin_and_gain_loss_code(pdf_path)

    WordDocExport(ThereturnDataTuple[0],ThereturnDataTuple[1])
