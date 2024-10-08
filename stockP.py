#Finds the payers for each stock item

import pdfplumber
from collections import defaultdict
import docx

#User Settings: Name of the script
TargetPDFFile = "Wage and Income Transcript.pdf"
OutPutDocFileName = "StockPayerWageIncTranscriptSummary"

def count_payer_fin_account(pdf_path):
    # Dictionary to store counts of each payer's FIN and name
    fin_name_counts = defaultdict(int)

    # Extract text using pdfplumber
    with pdfplumber.open(pdf_path) as pdf:
        total_text = ""
        for page in pdf.pages:
            total_text += page.extract_text()
    lines = total_text.split("\n")

    current_fin = None
    # Iterate through the lines to find relevant entries
    for i, line in enumerate(lines):
        if "Payer's Federal Identification Number (FIN):" in line:
            current_fin = line.split("Payer's Federal Identification Number (FIN):")[-1].strip()
            # Check the next line for the payer's name
            next_line = lines[i + 1]
            if next_line:
                name = next_line.split(',')[0].title()  # Convert to title case
                key = (current_fin, name)
                fin_name_counts[key] += 1

    return fin_name_counts

    

def WordDocExport(APayerdict):
    TheWdDoc = docx.Document()

    TheWdDoc.add_heading("Summary of Stock Payer in Transcript")

    
    for (fin, name), count in APayerdict.items():
        TheWdDoc.add_paragraph(f"Payer's Federal Identification Number (FIN): {fin} | Name: {name} -> Count: {count}")

    TheWdDoc.save(OutPutDocFileName + ".docx")

if __name__ == "__main__":
    pdf_path = r"Wage and Income Transcript.pdf"
    ReturnPayerDict = count_payer_fin_account(pdf_path)

    WordDocExport(ReturnPayerDict)
